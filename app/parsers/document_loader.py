from pathlib import Path

import pypdf
from docx import Document as WordDocument


class DocumentLoader:

    @staticmethod
    def load_txt(filepath):

        with open(filepath, "r", encoding="utf-8") as file:
            return file.read()

    @staticmethod
    def load_docx(filepath):

        document = WordDocument(filepath)

        text_blocks = []

        for paragraph in document.paragraphs:

            text = paragraph.text.strip()

            if text:
                text_blocks.append(text)

        for table in document.tables:

            for row in table.rows:

                cells = []

                for cell in row.cells:

                    value = cell.text.strip()

                    if value:
                        cells.append(value)

                if cells:
                    text_blocks.append(" | ".join(cells))

        return "\n".join(text_blocks)

    @staticmethod
    def load_pdf(filepath):

        reader = pypdf.PdfReader(filepath)

        pages = []

        for page in reader.pages:

            text = page.extract_text()

            if text:
                pages.append(text)

        return "\n".join(pages)

    @staticmethod
    def load(filepath):

        suffix = Path(filepath).suffix.lower()

        if suffix == ".docx":
            return DocumentLoader.load_docx(filepath)

        if suffix == ".pdf":
            return DocumentLoader.load_pdf(filepath)

        if suffix == ".txt":
            return DocumentLoader.load_txt(filepath)

        raise ValueError(f"Unsupported document type: {suffix}")