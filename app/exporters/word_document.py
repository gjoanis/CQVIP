from docx import Document
from datetime import date

from app.config.project_config import ProjectConfig
from app.services.document_control import DocumentControl


class WordDocumentExporter:

    def __init__(self, title, project_name, asset_name, document_type):
        self.title = title
        self.project_name = project_name
        self.asset_name = asset_name
        self.document_type = document_type
        self.document = Document()

        config = ProjectConfig()
        control = DocumentControl(config)
        self.document_control = control.build(document_type)

    def add_title_page(self):
        self.document.add_heading(self.title, level=0)

        table = self.document.add_table(rows=9, cols=2)
        table.style = "Table Grid"

        rows = [
            ("Project", self.document_control["project"]),
            ("Asset", self.asset_name),
            ("Document Type", self.document_type),
            ("Document Number", self.document_control["document_number"]),
            ("Revision", self.document_control["revision"]),
            ("Status", self.document_control["status"]),
            ("Prepared By", self.document_control["prepared_by"]),
            ("Approved By", self.document_control["approved_by"]),
            ("Date", self.document_control["date"])
        ]

        for row, values in zip(table.rows, rows):
            row.cells[0].text = values[0]
            row.cells[1].text = values[1]

        self.document.add_page_break()

    def add_document_information(self):
        self.document.add_heading("Document Information", level=1)

        table = self.document.add_table(rows=8, cols=2)
        table.style = "Table Grid"

        rows = [
            ("Document Title", self.title),
            ("Project", self.document_control["project"]),
            ("Asset", self.asset_name),
            ("Document Number", self.document_control["document_number"]),
            ("Revision", self.document_control["revision"]),
            ("Status", self.document_control["status"]),
            ("Document Type", self.document_type),
            ("Date", str(date.today()))
        ]

        for row, values in zip(table.rows, rows):
            row.cells[0].text = values[0]
            row.cells[1].text = values[1]

    def add_section_heading(self, heading):
        self.document.add_heading(heading, level=1)

    def add_paragraph(self, text):
        self.document.add_paragraph(text)

    def add_approval_table(self):
        self.document.add_heading("Approval Signatures", level=1)

        table = self.document.add_table(rows=5, cols=4)
        table.style = "Table Grid"

        headers = ["Role", "Name", "Signature", "Date"]
        roles = [
            "Prepared By",
            "Reviewed By",
            "Quality Assurance",
            "System Owner"
        ]

        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = header

        for row_index, role in enumerate(roles, start=1):
            table.rows[row_index].cells[0].text = role

    def add_responsibilities_table(self):
        self.document.add_heading("Responsibilities", level=1)

        table = self.document.add_table(rows=5, cols=2)
        table.style = "Table Grid"

        rows = [
            ("Role", "Responsibility"),
            ("Validation", "Execute testing and document results."),
            ("Engineering", "Provide technical support and system knowledge."),
            ("Quality Assurance", "Review and approve protocol execution."),
            ("Operations", "Support operation of the equipment during testing.")
        ]

        for row, values in zip(table.rows, rows):
            row.cells[0].text = values[0]
            row.cells[1].text = values[1]

    def add_prerequisites(self):
        self.document.add_heading("Prerequisites", level=1)

        items = [
            "Approved requirements are available.",
            "Applicable procedures are available.",
            "Required personnel are trained.",
            "Required equipment and instruments are available.",
            "Safety requirements have been reviewed."
        ]

        for item in items:
            self.document.add_paragraph(f"☐ {item}")

    def add_equipment_required(self):
        self.document.add_heading("Equipment / Instruments Required", level=1)

        table = self.document.add_table(rows=1, cols=3)
        table.style = "Table Grid"

        headers = [
            "Equipment / Instrument",
            "ID",
            "Calibration Required"
        ]

        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = header

        equipment = [
            ("Calibrated Instrument", "", "Yes"),
            ("Data Logger / Recorder", "", "As Applicable"),
            ("Computer / Workstation", "", "No"),
            ("Approved Procedure", "", "No")
        ]

        for item in equipment:
            row = table.add_row().cells
            row[0].text = item[0]
            row[1].text = item[1]
            row[2].text = item[2]

    def add_safety_precautions(self):
        self.document.add_heading("Safety Precautions", level=1)

        items = [
            "Follow applicable safety procedures.",
            "Wear required PPE.",
            "Do not bypass safety interlocks.",
            "Report deviations or unsafe conditions immediately."
        ]

        for item in items:
            self.document.add_paragraph(f"• {item}")

    def add_test_execution_table(self, procedure_steps):
        table = self.document.add_table(rows=1, cols=5)
        table.style = "Table Grid"

        headers = [
            "Step",
            "Action",
            "Expected Result",
            "Actual Result",
            "Pass / Fail"
        ]

        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = header

        for index, step in enumerate(procedure_steps, start=1):
            row = table.add_row().cells
            row[0].text = str(index)
            row[1].text = step
            row[2].text = "Result meets acceptance criteria."
            row[3].text = ""
            row[4].text = ""

    def save(self, filename):
        self.document.save(filename)
        print(f"Word document saved as {filename}")